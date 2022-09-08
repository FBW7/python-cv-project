from docx import Document
from docx.shared import Inches
import pyttsx3

document = Document()

def speak(text):
    pyttsx3.speak(text)


# profile picture
document.add_picture(
    "me.png", 
    width = Inches(1.5)
)

# name, phone number and email details
speak("What is your name?")
name = input("What is your name?\n")
phone_number = input("Kindly input your phone number:\n")
email = input("What is your email address?\n")

document.add_paragraph(
    name + " | " + phone_number + " | " + email
)

# about me
document.add_heading("About me")
document.add_paragraph(input("Tell us about yourself: "))

# work experience

more_experience = True
document.add_heading("Work Experience")

while more_experience:

    
    p = document.add_paragraph()

    company = input("Enter Company: ")
    position = input("Enter your position: ")
    from_date = input("Enter start date: ")
    to_date = input("Enter to date: ")

    p.add_run(company + " ").bold = True
    p.add_run(position + "\n").italic = True
    p.add_run(from_date + " - " + to_date + "\n").italic = True

    experience_details = input("Tell us about your experience at " + company + ": " )
    p.add_run(experience_details)
    m = input("Do you have more work experience? use 'y' or 'n' ")
    if m.lower() == 'y':
        more_experience
    elif m.lower() == 'n':
        more_experience = False
    else:
        break

document.add_heading("Training and Certification")
more_training = True

while more_training:
    t_c = document.add_paragraph()
    certificates = input("Give a list of your Training and Certification ")
    t_c.style = "List Bullet"
    t_c.add_run(certificates)
    t = input("Do you have more work training and certification? use 'y' or 'n' ")
    if t.lower() == 'y':
        more_training
    elif t.lower() == 'n':
        more_training = False
    else:
        break

# skills

document.add_heading("Skills")
while True:
    s = document.add_paragraph()

    skill = input("Add your skills here for better offer: ")
    s.style = "List Bullet"
    s.add_run(skill)

    more_skill = input("Do you have more skill? 'y' or 'n' ")
    if more_skill.lower() == "y":
        True
    else:
        break

# footer
section = document.sections[0]
footer = section.footer
f = footer.paragraphs[0]
f.text = "CV generated using FBW code"




document.save("cv.docx")




